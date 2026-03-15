# input/docx_parser.py
from docx import Document


def docx_to_text(docx_path: str) -> str:
    """
    提取 .docx 文本：段落 + 表格
    """
    doc = Document(docx_path)
    lines = []

    # 段落
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)

    # 表格
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                ct = (cell.text or "").strip()
                ct = ct.replace("\n", " ").strip()
                if ct:
                    row_text.append(ct)
            if row_text:
                lines.append(" | ".join(row_text))

    return "\n".join(lines).strip()